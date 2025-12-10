"""DOCX export functionality for proposals."""
from typing import Optional, Dict, Any
from io import BytesIO
import re


def export_to_docx(
    proposal_html: str,
    project_title: Optional[str] = None,
    sender_name: Optional[str] = None,
    submission_date: Optional[str] = None
) -> BytesIO:
    """
    Convert proposal HTML to DOCX format.

    Args:
        proposal_html: HTML content of the proposal
        project_title: Title of the project
        sender_name: Name of the sender
        submission_date: Date of submission

    Returns:
        BytesIO object containing the DOCX file
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX export. "
            "Install it with: pip install python-docx"
        )

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Parse HTML and convert to DOCX
    _parse_html_to_docx(doc, proposal_html)

    # Save to BytesIO
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer


def _parse_html_to_docx(doc, html_content: str):
    """Parse HTML content and add it to the DOCX document."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Remove HTML comments and script tags
    html_content = re.sub(r'<!--.*?-->', '', html_content, flags=re.DOTALL)
    html_content = re.sub(r'<script.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)

    # Split by major HTML tags
    parts = re.split(r'(<h1[^>]*>.*?</h1>|<h2[^>]*>.*?</h2>|<p[^>]*>.*?</p>|<ul[^>]*>.*?</ul>|<li[^>]*>.*?</li>|<strong>.*?</strong>|<em>.*?</em>)', html_content, flags=re.DOTALL | re.IGNORECASE)

    in_list = False
    current_paragraph = None

    for part in parts:
        if not part.strip():
            continue

        part = part.strip()

        # Handle H1 tags (cover page title)
        if re.match(r'<h1[^>]*>.*?</h1>', part, re.DOTALL | re.IGNORECASE):
            text = re.sub(r'<[^>]+>', '', part)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text.strip())
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.name = 'Calibri'
            doc.add_paragraph()  # Add spacing

        # Handle H2 tags (section headings)
        elif re.match(r'<h2[^>]*>.*?</h2>', part, re.DOTALL | re.IGNORECASE):
            # Close any open list
            if in_list:
                in_list = False

            text = re.sub(r'<[^>]+>', '', part)
            p = doc.add_paragraph()
            run = p.add_run(text.strip())
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = 'Calibri'
            doc.add_paragraph()  # Add spacing

        # Handle UL tags (lists)
        elif re.match(r'<ul[^>]*>', part, re.IGNORECASE):
            in_list = True

        # Handle LI tags (list items)
        elif re.match(r'<li[^>]*>', part, re.IGNORECASE):
            text = re.sub(r'<[^>]+>', '', part)
            # Remove leading dash if present
            text = re.sub(r'^-\s*', '', text.strip())

            # Apply formatting (bold, italic)
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)

        # Handle P tags (paragraphs)
        elif re.match(r'<p[^>]*>', part, re.IGNORECASE):
            # Close any open list
            if in_list:
                in_list = False

            text = re.sub(r'<[^>]+>', '', part)
            if text.strip():
                p = doc.add_paragraph()
                _add_formatted_text(p, text.strip())
                doc.add_paragraph()  # Add spacing

        # Handle standalone text (might be between tags)
        elif not part.startswith('<'):
            if part.strip():
                if in_list:
                    p = doc.add_paragraph(style='List Bullet')
                    _add_formatted_text(p, part.strip())
                else:
                    p = doc.add_paragraph()
                    _add_formatted_text(p, part.strip())

    # Ensure proper spacing at the end
    doc.add_paragraph()


def _add_formatted_text(paragraph, text: str):
    """Add text to paragraph with formatting (bold, italic)."""
    from docx.shared import Pt

    # Split by formatting tags
    parts = re.split(r'(<strong>.*?</strong>|<em>.*?</em>)', text, flags=re.DOTALL | re.IGNORECASE)

    for part in parts:
        if not part.strip():
            continue

        # Handle bold
        if re.match(r'<strong>.*?</strong>', part, re.DOTALL | re.IGNORECASE):
            bold_text = re.sub(r'<[^>]+>', '', part)
            run = paragraph.add_run(bold_text)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

        # Handle italic
        elif re.match(r'<em>.*?</em>', part, re.DOTALL | re.IGNORECASE):
            italic_text = re.sub(r'<[^>]+>', '', part)
            run = paragraph.add_run(italic_text)
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

        # Regular text
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

